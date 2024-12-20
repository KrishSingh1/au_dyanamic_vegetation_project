# -*- coding: utf-8 -*-
"""
Created on Thu Apr  4 17:58:38 2024

@author: krish
"""

import numpy as np
import pandas as pd
import geopandas as gpd
import matplotlib.pyplot as plt
from sklearn.model_selection import TimeSeriesSplit
from sklearn.model_selection import cross_validate
import graphviz 
import os 

# To Create a Report from: 
from docx import Document
from docx.shared import Inches

import sys
sys.path.append('/Users/krish/Desktop/DYNAMIC MODEL VEGETATION PROJECT/au_dyanamic_vegetation_project/STEP9_DATA_MODELLING_AND_EXPLORATION')

# Attempt intel optimisation patch
from sklearnex import patch_sklearn, unpatch_sklearn

from sklearn.ensemble import RandomForestRegressor
import random
from sklearn.metrics import mean_squared_error
from sklearn.pipeline import Pipeline
from PreprocessData import * # import from custom transformers
from sklearn.model_selection import GridSearchCV
from sklearn.inspection import PartialDependenceDisplay
from sklearn.inspection import permutation_importance
from skopt import BayesSearchCV
from sklearn.model_selection import KFold
from sklearn.model_selection import GroupKFold

from timeit import default_timer as timer

import lightgbm as lgb
import xgboost as xgb
from sklearn import tree


#%% Functions 

def plotPredictions(actual, prediction, TARGET, directory_plot_output, msg = '', split = '', fire_split = ''):
    fig, ax = plt.subplots(nrows = 3, figsize = (15,10))
    fig.suptitle(msg, fontsize=30)
    for i,v in enumerate(TARGET):
        actual[v].plot(ax=ax[i], color = 'blue', alpha = 0.4, linestyle='dashed', label = f'Observed {v.split("_")[0]}')
        prediction[f'prediction_{v}'].plot(ax=ax[i], color = 'orange', ylim = (0,100), label = f'Modelled {v.split("_")[0]}')
        ax[i].legend()
        ax[i].grid(True)
        if split:
            for s in split:
                ax[i].axvline(s, color='black', ls='--')
        if fire_split:
            for f in fire_split:
                ax[i].axvline(f, color='red', ls='--')
                
    plt.savefig(fname = directory_plot_output)
    plt.close()
#%% Main 


# sites_list = np.unique(['TCATCH0004', 'TCATNM0001', 'TCATNM0003', 'TCATCH0006'])

# Development (after 2015 set)
after_2015_sites = ['WAAPIL0003', 'NSABHC0023', 'TCATCH0006',
                 'WAAGAS0002', 'NSAMDD0014', 'NTAGFU0021', 
                 'NSANSS0001', 'SATSTP0005', 'QDASSD0015', 
                 'NTAFIN0002', 'NSANAN0002', 'QDAEIU0010'] # smalller subset 


tree_sites = ['NSABBS0001', 'NSACOP0001', 'NSAMDD0011', 'NSAMDD0020',
       'NSAMUL0003', 'NSANAN0001', 'NSANAN0002', 'NSANSS0001',
       'NSANSS0002', 'NSTSYB0003', 'NSTSYB0006', 'NTAFIN0003',
       'NTAFIN0015', 'NTAGFU0030', 'NTAGFU0034', 'QDABBN0002',
       'QDACHC0003', 'QDACYP0020', 'QDAGUP0006', 'QDAMGD0025',
       'QDAMUL0002', 'QDAMUL0003', 'QDASEQ0004', 'SAAFLB0008',
       'SATFLB0003', 'SATFLB0020', 'SATFLB0022', 'TCATCH0004',
       'WAACAR0002', 'WAAGAS0001', 'WAAPIL0023']

shrub_sites = ['NSABHC0011', 'NSAMDD0028', 'NSTSYB0005', 'NTAFIN0018',
       'QDABBS0010', 'QDACYP0018', 'QDAGUP0021', 'SAAEYB0021',
       'SAAEYB0028', 'SAAFLB0005', 'SAAGAW0008', 'SAAKAN0009',
       'SAASTP0023', 'SAASTP0033', 'SAASTP0034', 'SASMDD0009',
       'SASMDD0014', 'SATFLB0023', 'WAACAR0004', 'WAACOO0007',
       'WAACOO0016', 'WAACOO0024', 'WAACOO0026', 'WAACOO0027',
       'WAACOO0029', 'WAAGES0001', 'WAALSD0002', 'WAANUL0003',
       'WAAPIL0010']

grass_sites = ['NSABHC0023', 'NSAMDD0001', 'NSAMDD0014', 'NTADAC0001',
       'NTADMR0001', 'NTAFIN0002', 'NTAFIN0006', 'NTAGFU0014',
       'NTAGFU0020', 'NTAGFU0021', 'NTASTU0004', 'NTTDMR0003',
       'QDABBS0002', 'QDACYP0006', 'QDACYP0022', 'QDAEIU0005',
       'QDAEIU0010', 'QDAGUP0009', 'QDAGUP0019', 'QDAMGD0002',
       'QDAMGD0023', 'QDAMGD0024', 'QDASSD0015', 'SAAEYB0029',
       'SAAFLB0003', 'SASMDD0005', 'SATFLB0019', 'SATSTP0005',
       'TCATCH0006', 'TCATNM0001', 'TCATNM0003', 'VCAAUA0012',
       'WAAAVW0006', 'WAACOO0030', 'WAAGAS0002', 'WAANOK0006',
       'WAAPIL0003', 'WAAPIL0024', 'WAAPIL0031']


#smaller_subset = pd.read_csv('../DATASETS/Sites_Subset_20231010/ausplots_site_info/sites_subset.csv').copy()
#bigger_subset = pd.read_csv('../DATASETS/Sites_Bigger_Subset_20240124/ausplots_bigger_subset.csv').copy()
#smaller_list = np.unique(list(np.unique(smaller_subset.site_location_name.values)))

super_group_list = ['Desert Chenopod', 'Desert Forb', 'Desert Hummock.grass',
       'Desert Shrub', 'Desert Tree.Palm', 'Desert Tussock.grass',
       'Temp/Med Shrub', 'Temp/Med Tree.Palm', 'Temp/Med Tussock.grass',
       'Tropical/Savanna Tree.Palm', 'Tropical/Savanna Tussock.grass']
selected_super_group = 'Desert Tree.Palm'
super_groups_classified = pd.read_csv('C:/Users/krish/Desktop/DYNAMIC MODEL VEGETATION PROJECT/au_dyanamic_vegetation_project/DATASETS/AusPlots_Extracted_Data/Final/sites_super_classified.csv')
selected_super_group_list = super_groups_classified.loc[super_groups_classified['super_group'] == selected_super_group]['site_location_name']

sites_list = selected_super_group_list
results_dir = f'C:/Users/krish/Desktop/DYNAMIC MODEL VEGETATION PROJECT/au_dyanamic_vegetation_project/RESULTS/Random_Forest_Results_On_Super_Group_Results/{selected_super_group}'


#%% Model the dataset

SEASONAL_FEATURES = ['photoperiod', 'photoperiod_gradient']

PRECIP_FEATURES = ['precip_30', 'precip_90', 'precip_180', 
                   'precip_365', 'precip_730', 'precip_1095', 
                   'precip_1460', 'MAP']

TEMP_FEATURES = ['tmax_lag', 'tmax_7', 'tmax_14', 
                 'tmax_30', 'tmin_lag', 'tmin_7', 
                 'tmin_14', 'tmin_30', 'MAT']

VPD_FEATURES = ['VPD_lag','VPD_7', 'VPD_14',
                'VPD_30']

LAG_FEATURES = ['pv_lag', 'npv_lag', 'bs_lag']

LAGGED_CHANGE_FEATURES = ['pv_change', 'npv_change', 'bs_change']

FIRE_FEATURES = ['days_since_fire', 'fire_severity']

CO2_FEATURES = ['CO2']

VEGETATION_FEATURES = ['grass', 'shrub', 'tree']

SOIL_FEATURES = ['SLGA_1','SLGA_2','SLGA_3', 'DER_000_999'] # the soil attributes to include

FEATURES =  SEASONAL_FEATURES + PRECIP_FEATURES + TEMP_FEATURES + VPD_FEATURES + FIRE_FEATURES + CO2_FEATURES + VEGETATION_FEATURES + SOIL_FEATURES # final features 
TARGET = ['pv_filter', 'npv_filter', 'bs_filter']
scores = []


#%% Create Train/test set 


# Training and test set 
datasets =  {} # entire set - for final evaluation 
training_set = {} # training set 
test_set = {} # test set 

# Iterate through the site list 

# Set up a random seed, based on the date it was set YYYYMMDD
random.seed(20240514)
number_of_blocks = 5
choices = [b for b in range(number_of_blocks)]
number_of_choices = len(sites_list)/len(choices)
duplicator = [round(np.floor(number_of_choices)) for i in range(len(choices))]

already_chosen = []
while sum(duplicator) != len(sites_list): # if there is an uneven split, keep adding 1 more until it sums to the avaliable number of datasets 
    chosen_index = random.randrange(0,len(duplicator),1)
    if chosen_index not in already_chosen:
        duplicator[chosen_index] += 1
        already_chosen.append(chosen_index)
    
random.shuffle(duplicator) # In cases where there is an uneven split, randomise which block gets the extra choice 
print(duplicator)

choice_adj = []
for index ,i in enumerate(choices):
    for j in range(duplicator[index]):
        choice_adj.append(i)
        
random.shuffle(choice_adj)
print(choice_adj)

# Now Construct the training and test dataset 

for i, site_location_name in enumerate(sites_list):
    site_merged = pd.read_csv(f'../DATASETS/DEA_FC_PROCESSED/MODELLED_PREPROCESSED/Input_DataSet_{site_location_name}.csv',
                              parse_dates = ['time']).copy().dropna(subset = FEATURES) # read and drop na
    site_merged.sort_values('time', inplace = True)
    site_merged.reset_index(inplace = True)
    
    period = len(site_merged)//number_of_blocks # get size of time period (as expressed by the number of data points)
    lower_bound = choice_adj[i] * period
    if lower_bound == choices[-1] * period: # if its the last block%, simply take all time points from there up to the most recent time point
        upper_bound = len(site_merged) -1
    else:
        upper_bound = lower_bound + period 

    # get test set 
    test = site_merged[(site_merged.index >= lower_bound) & (site_merged.index <= upper_bound)]
    print(f'{site_location_name} Test Set: {(site_merged.time[lower_bound], site_merged.time[upper_bound])}, Period: {period}')
    # get train set (note the selection condition is logically opposite to selection condition of the test set )
    train = site_merged[(site_merged.index < lower_bound) | (site_merged.index > upper_bound)]
    
    datasets[site_location_name] = site_merged
    training_set[site_location_name] = train
    test_set[site_location_name] = test
    
training_merged = pd.concat(training_set).dropna(subset = FEATURES) # drop na based on chosen features, needed for random forest 
training_merged.sort_values('time', inplace = True)
training_merged.set_index('time', inplace = True)

test_merged = pd.concat(test_set).dropna(subset = FEATURES)
test_merged.sort_values('time', inplace = True)
test_merged.set_index('time', inplace = True)


#%% Determine Optimal Model HyperParams 

random_state = 20240228
main_scorer = 'neg_mean_squared_error'
# Possible scorers:
    #  neg_mean_absolute_percentage_error'
    #  mean_squared_log_error
    # See more below:
    # https://scikit-learn.org/stable/modules/model_evaluation.html#regression-metrics
    
## Test with Default RF 

criterion = ['squared_error', 'friedman_mse', 'poisson']
# Skip absolute error, as that is no optimised well 
max_depth = range(1, 56, 5)
oob_scores = []

start = timer()
for n_est in max_depth:
    print(f'Trying {n_est}')
    reg = RandomForestRegressor(max_depth= n_est, random_state = random_state, n_jobs = 8, oob_score= True)
    reg.fit(X = training_merged[FEATURES], y = training_merged[TARGET])
    oob_scores.append(reg.oob_score_)
end = timer()

plt.plot(max_depth, oob_scores)

#start = timer()
#reg = RandomForestRegressor(n_estimators = 100, random_state = random_state, n_jobs = 8)
#reg.fit(X = training_merged[FEATURES], y = training_merged[TARGET])
#end = timer()
print(f' Time taken: {end - start}')

print(([estimator.tree_.max_depth for estimator in reg.estimators_]))
# The maximum tree depth is interestingly 41 
     

#fig, axes = plt.subplots(nrows = 1, ncols = 1, figsize = (4,4), dpi=400)
#tree_number = 0
#tree.plot_tree(reg.estimators_[tree_number],
#               feature_names = FEATURES, 
#               class_names= TARGET,
#               filled = True);
#fig.savefig(f'{results_dir}/Random_Forest_{selected_super_group}_{tree_number}.png')

#print(f' Time taken: {end - start}')

#plt.plot(n_estimators, oob_scores)



# default_RF_CV = pd.DataFrame(cross_validate(reg, X = training_merged[FEATURES],
#                      y = training_merged[TARGET], cv=cv_splits, 
#                      scoring=np.unique(['neg_mean_squared_error', main_scorer]).tolist(),
#                      return_train_score = True))
# print(f'Default\nTrain R2 {default_RF_CV["train_" + "neg_mean_squared_error"].mean()}\nTest R2 {default_RF_CV["test_" + "neg_mean_squared_error"].mean()}')


#%% Fit the Models

# Default Model 

print(mean_squared_error(reg.predict(training_merged[FEATURES]), training_merged[TARGET]))
print(mean_squared_error(reg.predict(training_merged[FEATURES]), training_merged[TARGET], multioutput = 'raw_values'))
print(mean_squared_error(reg.predict(test_merged[FEATURES]), test_merged[TARGET]))
print(mean_squared_error(reg.predict(test_merged[FEATURES]), test_merged[TARGET], multioutput = 'raw_values'))

historical_fire_ds = pd.read_csv('../DATASETS/AusPlotsBurnData/Combined_Data/AusPlots_Combined_Fire_Dataset.csv', parse_dates = ['ignition_d']) # Fire Dataset

#%% Time To get results and generate a report for visualisation 

results_dir = f'C:/Users/krish/Desktop/DYNAMIC MODEL VEGETATION PROJECT/au_dyanamic_vegetation_project/RESULTS/Random_Forest_Results_On_Super_Group_Results/{selected_super_group}'

# Create Directory If this does not exist 
if os.path.exists(results_dir) == False:
    os.makedirs(results_dir)
    os.makedirs(results_dir + '/Plots') # also create a plots folder under since we can assume this folder did not initialy exist 
    
doc = Document()

for site in sites_list:

    site_data = datasets[site].set_index('time').dropna(subset = FEATURES)
    train_site = site_data.iloc[training_set[site].index]
    test_site = site_data.iloc[test_set[site].index]

    print(f'{test_site.index.min()}')
    
    y_pred = reg.predict(site_data[FEATURES])
    TARGET_names = ['prediction_' + i for i in TARGET]
    df = pd.DataFrame(y_pred, columns = TARGET_names)
    df.index = site_data[FEATURES].index
    
    train_score = mean_squared_error(train_site[TARGET], df.iloc[training_set[site].index])
    test_score = mean_squared_error(test_site[TARGET],  df.iloc[test_set[site].index])
    
    historical_fire_pipeline = Pipeline([
        ('historical_burn_date_preprocess', historical_burn_date_preprocess(site))
        ])
    historical_fire_ds_site = historical_fire_pipeline.fit_transform(historical_fire_ds)
    
    dates = ''  # make dates null unless there are fire dates
    if historical_fire_ds_site.empty == False:
        dates = [i for i in historical_fire_ds_site['ignition_d'] if pd.isnull(i) == False] # check if the record does not have null values, otherwise filter it out
        
    plotPredictions(site_data, df, TARGET, msg = f'{site} MSE Scores: train:{train_score:.2f}, test:{test_score:.2f}',
                    split = [f'{test_site.index.min()},', f'{test_site.index.max()}'],
                    fire_split = dates, directory_plot_output = f"{results_dir}/Plots/{site}.png")
    
    doc.add_heading(f'{site}', level=1)
    doc.add_picture(f"{results_dir}/Plots/{site}.png", width=Inches(7))  
    doc.add_page_break() 
    doc.save(f'{results_dir}/Random_Forest_{selected_super_group}_Results.docx')



#%% Examine Importances 

# Get importance 
# https://scikit-learn.org/stable/modules/generated/sklearn.inspection.permutation_importance.html#sklearn.inspection.permutation_importance
# https://scikit-learn.org/stable/auto_examples/inspection/plot_permutation_importance_multicollinear.html#sphx-glr-auto-examples-inspection-plot-permutation-importance-multicollinear-py
# Comparion between plots code inspired by above links 

#reg = RandomForestRegressor(**grid.best_params_)
# reg.fit(training_merged[FEATURES], training_merged[TARGET])

# # fig, ax = plt.subplots(1,2)
# gini_importance = pd.DataFrame(reg.feature_importances_.T, index = FEATURES, columns = ['Gini_importance'])
# gini_importance.sort_values(by = 'Gini_importance', inplace = True, ascending = True)
# gini_importance.plot.barh(figsize = (15,5))
# #gini_importance.to_csv('Gini_Importance_trees.csv')


n_repeats = 50
perm_importance = permutation_importance(reg, test_merged[FEATURES], test_merged[TARGET],
                                          n_repeats=n_repeats, random_state = random_state, n_jobs = 7, scoring = 'neg_mean_squared_error')
perm_sorted_idx = perm_importance.importances_mean.argsort()
perm_importance_df = pd.DataFrame(perm_importance.importances[perm_sorted_idx].T, columns = test_merged[FEATURES].columns[perm_sorted_idx])
# perm_importance_df.boxplot(vert = False, figsize = (15,10))
# fig.tight_layout()

arr_importances = np.array([list(perm_importance['importances_mean']), list(perm_importance['importances_std'])]).T
perm_importance_df_2 = pd.DataFrame(arr_importances, columns = ['importances_mean', 'importances_std'], index = FEATURES)
perm_importance_df_2.sort_values('importances_mean', ascending = True, inplace = True)
perm_importance_df_2.plot.barh(figsize = (15, 10))
# #perm_importance_df_2.to_csv('Perm_importance_trees.csv')
