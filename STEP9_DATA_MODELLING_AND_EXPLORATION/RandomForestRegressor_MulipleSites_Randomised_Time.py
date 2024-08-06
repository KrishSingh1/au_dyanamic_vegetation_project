# -*- coding: utf-8 -*-
"""
Created on Thu Apr  4 17:58:38 2024

@author: krish
"""

import numpy as np
import pandas as pd
import geopandas as gpd
import matplotlib.pyplot as plt
from sklearn.model_selection import TimeSeriesSplit
from sklearn.model_selection import cross_validate
import graphviz 
import os 

# To Create a Report from: 
from docx import Document
from docx.shared import Inches

import sys
sys.path.append('/Users/krish/Desktop/DYNAMIC MODEL VEGETATION PROJECT/au_dyanamic_vegetation_project/STEP9_DATA_MODELLING_AND_EXPLORATION')

from sklearn.ensemble import RandomForestRegressor
import random
from sklearn.metrics import mean_squared_error
from sklearn.pipeline import Pipeline
from PreprocessData import * # import from custom transformers
from sklearn.model_selection import GridSearchCV
from sklearn.inspection import PartialDependenceDisplay
from sklearn.inspection import permutation_importance
from skopt import BayesSearchCV
from sklearn.model_selection import KFold
from sklearn.model_selection import GroupKFold

import alibi
from alibi.explainers import ALE, plot_ale

#%% Functions 

def plotPredictions(actual, prediction, TARGET, directory_plot_output, msg = '', split = '', fire_split = ''):
    fig, ax = plt.subplots(nrows = 3, figsize = (15,10))
    fig.suptitle(msg, fontsize=30)
    for i,v in enumerate(TARGET):
        actual[v].plot(ax=ax[i], color = 'blue', alpha = 0.4, linestyle='dashed', label = f'Observed {v.split("_")[0]}')
        prediction[f'prediction_{v}'].plot(ax=ax[i], color = 'orange', ylim = (0,100), label = f'Modelled {v.split("_")[0]}')
        ax[i].legend()
        ax[i].grid(True)
        if split:
            for s in split:
                ax[i].axvline(s, color='black', ls='--')
        if fire_split:
            for f in fire_split:
                ax[i].axvline(f, color='red', ls='--')
                
    plt.savefig(fname = directory_plot_output)
    plt.close()
#%% Main 

super_group_list = ['Desert Chenopod', 'Desert Forb', 'Desert Hummock.grass',
       'Desert Shrub', 'Desert Tree.Palm', 'Desert Tussock.grass',
       'Temp/Med Shrub', 'Temp/Med Tree.Palm', 'Temp/Med Tussock.grass',
       'Tropical/Savanna Tree.Palm', 'Tropical/Savanna Tussock.grass']

target_group_list = [] 
for s in super_group_list:
    print(s)
    
    selected_super_group = s
    super_groups_classified = pd.read_csv('C:/Users/krish/Desktop/DYNAMIC MODEL VEGETATION PROJECT/au_dyanamic_vegetation_project/DATASETS/AusPlots_Extracted_Data/Final/sites_super_classified.csv')
    selected_super_group_list = super_groups_classified.loc[super_groups_classified['super_group'] == selected_super_group]['site_location_name']
    
    sites_list = selected_super_group_list
    
    
    #%% Model the dataset
    
    SEASONAL_FEATURES = ['photoperiod', 'photoperiod_gradient']
    
    PRECIP_FEATURES = ['precip_30', 'precip_90', 'precip_180', 
                       'precip_365', 'precip_730', 'precip_1095', 
                       'precip_1460', 'MAP']
    
    TEMP_FEATURES = ['tmax_lag', 'tmax_7', 'tmax_14', 
                     'tmax_30', 'tmin_lag', 'tmin_7', 
                     'tmin_14', 'tmin_30', 'MAT']
    
    VPD_FEATURES = ['VPD_lag','VPD_7', 'VPD_14',
                    'VPD_30']
    
    LAG_FEATURES = ['pv_lag', 'npv_lag', 'bs_lag']
    
    LAGGED_CHANGE_FEATURES = ['pv_change', 'npv_change', 'bs_change']
    
    FIRE_FEATURES = ['days_since_fire', 'fire_severity']
    
    CO2_FEATURES = ['CO2']
    
    VEGETATION_FEATURES = ['grass', 'shrub', 'tree']
    
    SOIL_FEATURES = ['SLGA_1','SLGA_2','SLGA_3', 'DER_000_999'] # the soil attributes to include
    
    FEATURES =  SEASONAL_FEATURES + PRECIP_FEATURES + TEMP_FEATURES + VPD_FEATURES + FIRE_FEATURES + CO2_FEATURES + VEGETATION_FEATURES + SOIL_FEATURES # final features 
    TARGET = ['pv_filter', 'npv_filter', 'bs_filter']
    scores = []
    
    
    #%% Create Train/test set 
    
    
    # Training and test set 
    datasets =  {} # entire set - for final evaluation 
    training_set = {} # training set 
    test_set = {} # test set 
    
    # Iterate through the site list 
    # Set up a random seed, based on the date it was set YYYYMMDD
    # 
    random.seed(20240514)
    
    # Now Construct the training and test dataset 
    to_print_as_word = []
    for i, site_location_name in enumerate(sites_list):
        site_merged = pd.read_csv(f'../DATASETS/DEA_FC_PROCESSED/MODELLED_PREPROCESSED/Input_DataSet_{site_location_name}.csv',
                                  parse_dates = ['time']).copy().dropna(subset = FEATURES) # read and drop na
        site_merged.sort_values('time', inplace = True)
        site_merged.reset_index(inplace = True)
        
        datasets[site_location_name] = site_merged
        training_set[site_location_name] = train
        test_set[site_location_name] = test
        
    training_merged = pd.concat(training_set).dropna(subset = FEATURES) # drop na based on chosen features, needed for random forest 
    training_merged.sort_values('time', inplace = True)
    training_merged.set_index('time', inplace = True)
    
    test_merged = pd.concat(test_set).dropna(subset = FEATURES)
    test_merged.sort_values('time', inplace = True)
    test_merged.set_index('time', inplace = True)
    
    
    
    #%% Run The model 
    
    random_state = 20240228
    main_scorer = 'neg_mean_squared_error'
    # Possible scorers:
        #  neg_mean_absolute_percentage_error'
        #  mean_squared_log_error
        # See more below:
        # https://scikit-learn.org/stable/modules/model_evaluation.html#regression-metrics
    
    reg = RandomForestRegressor(n_estimators = 100, random_state = random_state, n_jobs = 8)
    #%% Fit the Models
    
    reg.fit(X = training_merged[FEATURES], y = training_merged[TARGET])
    
    
    historical_fire_ds = pd.read_csv('../DATASETS/AusPlotsBurnData/Combined_Data/AusPlots_Combined_Fire_Dataset.csv', parse_dates = ['ignition_d']) # Fire Dataset
    
    training_row = mean_squared_error(reg.predict(training_merged[FEATURES]), training_merged[TARGET], multioutput = 'raw_values')
    testing_row = mean_squared_error(reg.predict(test_merged[FEATURES]), test_merged[TARGET], multioutput = 'raw_values')
    training_mean = mean_squared_error(reg.predict(training_merged[FEATURES]), training_merged[TARGET])
    testing_mean = mean_squared_error(reg.predict(test_merged[FEATURES]), test_merged[TARGET])
    results_table = pd.DataFrame([training_row, testing_row], columns = TARGET)
    results_table['loss_mean'] = [training_mean, testing_mean]
    results_table['training_set'] = ['Train', 'Test']
    results_table.set_index('training_set', inplace = True)
    
    print(results_table)
    #%% Time To get results and generate a report for visualisation 
    
    selected_super_group = '_'.join(s.split('/')) 
    results_dir = f'C:/Users/krish/Desktop/DYNAMIC MODEL VEGETATION PROJECT/au_dyanamic_vegetation_project/RESULTS/Random_Forest_Results_On_Super_Group_Results_2/{selected_super_group}'
    
    # Create Directory If this does not exist 
    if os.path.exists(results_dir) == False:
        os.makedirs(results_dir)
    
    if os.path.exists(results_dir + '/Plots') == False:
        os.makedirs(results_dir + '/Plots') # also create a plots folder under since we can assume this folder did not initialy exist 
    
    if os.path.exists(results_dir + '/Results') == False:
        os.makedirs(results_dir + '/Results')
        
    if os.path.exists(results_dir + '/Training') == False:
        os.makedirs(results_dir + '/Training')
    
    #%% Create CSV files 
    
    training_merged.to_csv(f'{results_dir}/Training/{selected_super_group}_Training_set.csv')
    test_merged.to_csv(f'{results_dir}/Training/{selected_super_group}_Test_set.csv')
    results_table.to_csv(f'{results_dir}/Results/{selected_super_group}_Overall_MSE.csv')
    
    print(f'Size of training and test set for {selected_super_group}')
    print(len(training_merged))
    print(len(test_merged))
    
    continue
    #%% Create document 
        
    doc = Document()
    
    doc.add_heading('Sites that were included', level=1)
    
    text = '\n\n'.join(to_print_as_word) 
    doc.add_paragraph(text)
    doc.add_paragraph(' ')
    
    doc.add_heading('Table: MSE of the Random Forest on each output', level=1)
    table_variable = doc.add_table(rows = results_table.shape[0] + 1,
                                    cols = results_table.shape[1] + 1)
        
    for i, row in enumerate(table_variable.rows):
        for j, cell in enumerate(row.cells):
            if i == 0:
                all_columns = list(results_table.columns) + ['']
                cell.text = all_columns[j-1]
            elif j == 0:
                cell.text = results_table.index[i-1]
            else:
                cell.text = str(round(results_table.iloc[i-1,j-1], 3))
    table_variable.style = 'Table Grid'
    
    doc.add_page_break() 
    
    
    
    #%% Examine Importances 
    
    # Get importance 
    # https://scikit-learn.org/stable/modules/generated/sklearn.inspection.permutation_importance.html#sklearn.inspection.permutation_importance
    # https://scikit-learn.org/stable/auto_examples/inspection/plot_permutation_importance_multicollinear.html#sphx-glr-auto-examples-inspection-plot-permutation-importance-multicollinear-py
    # Comparion between plots code inspired by above links 
    
    n_repeats = 50
    perm_importance = permutation_importance(reg, test_merged[FEATURES], test_merged[TARGET],
                                              n_repeats=n_repeats, random_state = random_state, n_jobs = 8, scoring = 'neg_mean_squared_error')
    perm_sorted_idx = perm_importance.importances_mean.argsort()
    perm_importance_df = pd.DataFrame(perm_importance.importances[perm_sorted_idx].T, columns = test_merged[FEATURES].columns[perm_sorted_idx])
    
    arr_importances = np.array([list(perm_importance['importances_mean']), list(perm_importance['importances_std'])]).T
    perm_importance_df_2 = pd.DataFrame(arr_importances, columns = ['importances_mean', 'importances_std'], index = FEATURES)
    perm_importance_df_2.sort_values('importances_mean', ascending = True, inplace = True)
    
    fig, ax = plt.subplots(1)
    
    perm_importance_df_2['importances_mean'].plot.barh(figsize = (15, 10), ax = ax)
    xticks = [i for i in range(0, round(max(perm_importance_df_2['importances_mean']) + 10), 10)]
    ax.set_xticks(xticks)
    plt.grid(True)
    plt.xlabel('Mean Gain in MSE')
    
    directory_plot_output = f"{results_dir}/Plots/Permutation_Importance.png"
    plt.savefig(fname = directory_plot_output)
    plt.close()
    
    doc.add_heading('Permutation Importance', level=1)
    doc.add_picture(directory_plot_output, width=Inches(7))  
    
    
    
    #%% Permutation Importance by Precip Vars
    
    precip_var = ['precip_30', 'precip_90', 'precip_180', 'precip_365', 'precip_730', 'precip_1095', 'precip_1460'][::-1]
    precip_var_set = ['1-30 days', '31-90 days', '91-180 days', '181-365 days', '366-730 days', '731-1095 days', '1096-1460 days'][::-1]
    
    fig, ax = plt.subplots(1, figsize = (10, 5))
    
    t = perm_importance_df_2.loc[precip_var]
    t['precip_var_set'] = precip_var_set
    t = t.set_index('precip_var_set')
    t.index.name = None
    t['importances_mean'].plot.barh(ax = ax)
    plt.grid(True)
    plt.xlabel('Mean Gain in MSE')
    
    directory_plot_output = f"{results_dir}/Plots/Permutation_Importance_Precip.png"
    plt.savefig(fname = directory_plot_output)
    plt.close()
    
    doc.add_heading('Permutation Importance by Precip', level=1)
    doc.add_picture(directory_plot_output, width=Inches(7))  
    doc.add_page_break() 
    
    
    #%% ALE plots
    
    rf_ale = ALE(reg.predict, feature_names=FEATURES, target_names=TARGET)
    rf_exp_tree = rf_ale.explain(np.array(training_merged[FEATURES]))
    
    nrow = 7
    ncol = 5
    fig, ax = plt.subplots(nrow, ncol, figsize = (15,20), sharey = True)
    
    # Iterate through row, then cols
    counter = 0
    for row in range(nrow):
        for col in range(ncol):
            # Break the loop when the number of features is limited 
            if counter == len(FEATURES):
                break
            f = FEATURES[counter]
            
            ax[row, col].plot(rf_exp_tree['data']['feature_values'][FEATURES.index(f)], 
                    rf_exp_tree['data']['ale_values'][FEATURES.index(f)][:, 0], 
                    color = 'green', label = 'PV', marker='o', linewidth= 0.5, markersize= 1.5)
            ax[row, col].plot(rf_exp_tree['data']['feature_values'][FEATURES.index(f)], 
                       rf_exp_tree['data']['ale_values'][FEATURES.index(f)][:, 1],
                    color = 'blue', label = 'NPV', marker='o', linewidth= 0.5, markersize= 1.5)
            ax[row, col].plot(rf_exp_tree['data']['feature_values'][FEATURES.index(f)], 
                       rf_exp_tree['data']['ale_values'][FEATURES.index(f)][:, 2],
                    color = 'brown', label = 'BS', marker='o', linewidth= 0.5, markersize= 1.5)
            ax[row, col].axhline(y=0, color='black', linestyle='--')
            ax[row, col].grid(True)
            ax[row, col].set_xlabel(f)
            
            if col == 0:
                ax[row, col].set_ylabel('ALE')
            counter += 1
    plt.tight_layout()
    
    directory_plot_output = f"{results_dir}/Plots/ALE_Plots.png"
    plt.savefig(fname = directory_plot_output)
    plt.close()
    
    doc.add_heading(f'ALE Importance', level=1)
    doc.add_picture(directory_plot_output, width=Inches(7))  
    doc.add_page_break() 
    
    
    #%% Plot Time series 
    
    for site in sites_list:
    
        site_data = datasets[site].set_index('time').dropna(subset = FEATURES)
        train_site = site_data.iloc[training_set[site].index]
        test_site = site_data.iloc[test_set[site].index]
    
        print(f'{test_site.index.min()}')
        
        y_pred = reg.predict(site_data[FEATURES])
        TARGET_names = ['prediction_' + i for i in TARGET]
        df = pd.DataFrame(y_pred, columns = TARGET_names)
        df.index = site_data[FEATURES].index
        
        train_score = mean_squared_error(train_site[TARGET], df.iloc[training_set[site].index])
        test_score = mean_squared_error(test_site[TARGET],  df.iloc[test_set[site].index])
        
        historical_fire_pipeline = Pipeline([
            ('historical_burn_date_preprocess', historical_burn_date_preprocess(site))
            ])
        historical_fire_ds_site = historical_fire_pipeline.fit_transform(historical_fire_ds)
        
        dates = ''  # make dates null unless there are fire dates
        if historical_fire_ds_site.empty == False:
            dates = [i for i in historical_fire_ds_site['ignition_d'] if pd.isnull(i) == False] # check if the record does not have null values, otherwise filter it out
            
        plotPredictions(site_data, df, TARGET, msg = f'{site} MSE Scores: train:{train_score:.2f}, test:{test_score:.2f}',
                        split = [f'{test_site.index.min()},', f'{test_site.index.max()}'],
                        fire_split = dates, directory_plot_output = f"{results_dir}/Plots/{site}.png")
        
        doc.add_heading(f'{site}', level=1)
        doc.add_picture(f"{results_dir}/Plots/{site}.png", width=Inches(7))  
        doc.add_page_break() 
        
    #%% Save the Document
    
    doc.save(f'{results_dir}/Random_Forest_{selected_super_group}_Results.docx')
